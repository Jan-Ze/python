from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
wb = load_workbook('./封号名单.xlsx')
ws = wb.active

for message in ws.iter_rows(min_row = 2,max_row = 26,values_only = True):
    name = message[0]
    wechat = message[1]
    doc = Document('./法务函模板.docx')
    para = doc.paragraphs[5]
    run1 = para.add_run()
    run1.add_text(name)
    run1.font.size = Pt(14)
    run1.font.underline = True
    run1.font.bold = True
    run2 = para.add_run()
    run2.add_text('同学 (WeChat ID: ' + wechat + ')')
    run2.font.size = Pt(14)
    doc.save('./new' + '\\' + '\\' + '法务函-' + name + '.docx')
    
    

    